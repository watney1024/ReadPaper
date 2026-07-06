"""
ReadPaper DOCX 阅读笔记生成器
用法：python generate_reading_notes.py <data.json路径> <output.docx路径>
"""
import json
import sys
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt


def add_labeled(doc, label, text):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)
    return p


def validate_data(data):
    required_top = [
        'title', 'keywords', 'abstract_en', 'abstract_zh',
        'intro_paragraphs', 'intro_logic', 'methods', 'results',
        'results_logic', 'disc_paragraphs', 'disc_logic',
    ]
    missing_top = [k for k in required_top if k not in data]
    if missing_top:
        raise ValueError(f"JSON data is missing required top-level keys: {missing_top}")

    required_keywords = ['research_question', 'techniques', 'paradigm', 'analysis_methods', 'conclusions', 'innovations']
    missing_kw = [k for k in required_keywords if k not in data['keywords']]
    if missing_kw:
        raise ValueError(f"JSON data['keywords'] is missing required sub-keys: {missing_kw}")

    required_methods = ['participants', 'paradigm', 'procedure', 'analysis']
    missing_m = [k for k in required_methods if k not in data['methods']]
    if missing_m:
        raise ValueError(f"JSON data['methods'] is missing required sub-keys: {missing_m}")


def create_reading_notes(data_path, output_path):
    try:
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"[ERROR] JSON file not found: {data_path}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"[ERROR] Failed to parse JSON file '{data_path}': {e}")
        sys.exit(1)

    validate_data(data)

    doc = Document()

    # ── 标题 ──
    doc.add_heading('论文阅读笔记', level=0)
    doc.add_heading(data['title'], level=1)
    doc.add_paragraph(f"生成日期：{data.get('date', datetime.now().strftime('%Y-%m-%d'))}")

    # ── 一、关键词卡片 ──
    doc.add_heading('一、关键词卡片', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '维度'
    hdr[1].text = '内容'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    kw = data['keywords']
    for i, (label, value) in enumerate([
        ('研究问题',     kw['research_question']),
        ('技术手段',     kw['techniques']),
        ('实验范式',     kw['paradigm']),
        ('数据分析方法', kw['analysis_methods']),
        ('主要结论',     kw['conclusions']),
        ('创新点',       kw['innovations']),
    ], start=1):
        row = table.rows[i].cells
        row[0].text = label
        row[1].text = value
    doc.add_paragraph()

    # ── 二、Abstract ──
    doc.add_heading('二、Abstract', level=2)
    doc.add_heading('原文（英文）', level=3)
    doc.add_paragraph(data['abstract_en'])
    doc.add_heading('中文翻译', level=3)
    doc.add_paragraph(data['abstract_zh'])

    # ── 三、Introduction ──
    doc.add_heading('三、Introduction 各段摘要与行文逻辑', level=2)
    for i, para in enumerate(data['intro_paragraphs'], start=1):
        doc.add_paragraph(f"【第{i}段 · 定位词：{para['locator']}】")
        add_labeled(doc, '▸ 主要内容：', para['summary'])
        if para.get('citations'):
            add_labeled(doc, '  引用文献：', para['citations'])
    doc.add_paragraph()
    add_labeled(doc, '▸ 行文逻辑总结：', data['intro_logic'])

    # ── 四、Methods ──
    doc.add_heading('四、Methods 摘要', level=2)
    m = data['methods']
    add_labeled(doc, '被试：',       m['participants'])
    add_labeled(doc, '实验范式：',   m['paradigm'])
    add_labeled(doc, '实验流程：',   m['procedure'])
    add_labeled(doc, '数据分析方法：', m['analysis'])

    # ── 五、Results ──
    doc.add_heading('五、Results 逐一对应', level=2)
    for i, item in enumerate(data['results'], start=1):
        add_labeled(doc, f'[{i}] 分析方法：', item['method'])
        add_labeled(doc, '→ 结果：',          item['result'])
        add_labeled(doc, '→ 结论：',          item['conclusion'])
        doc.add_paragraph()
    add_labeled(doc, '▸ 行文逻辑总结：', data['results_logic'])

    # ── 六、Discussion ──
    doc.add_heading('六、Discussion 各段摘要与行文逻辑', level=2)
    for i, para in enumerate(data['disc_paragraphs'], start=1):
        doc.add_paragraph(f"【第{i}段 · 定位词：{para['locator']}】")
        add_labeled(doc, '▸ 主要内容：', para['summary'])
    doc.add_paragraph()
    add_labeled(doc, '▸ 行文逻辑总结：', data['disc_logic'])

    # ── 七、我的阅读笔记 ──
    doc.add_heading('七、我的阅读笔记（待填写）', level=2)
    for name in ['核心贡献与意义', '与我研究的关联', '方法借鉴点', '疑问与待深入之处', '其他备注']:
        add_labeled(doc, f'【{name}】', '')
        doc.add_paragraph()
        doc.add_paragraph()

    output_dir = os.path.dirname(os.path.abspath(output_path))
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    print(f"[OK] DOCX saved: {output_path}")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("用法：python generate_reading_notes.py <data.json> <output.docx>")
        sys.exit(1)
    create_reading_notes(sys.argv[1], sys.argv[2])
